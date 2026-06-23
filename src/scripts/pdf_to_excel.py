"""
pdf_to_excel.py  –  PDF/Word→Excel 変換スクリプト
優先順位（PDF）:
  1. Azure Document Intelligence (prebuilt-layout) - 縦書き・複雑表・スキャン対応
  2. pdfplumber テーブル検出（線ありPDF・フォールバック）
  3. PyMuPDF テキスト抽出（フォールバック）
  4. pdfplumber テキスト抽出（最終フォールバック）
優先順位（Word .docx）:
  1. Azure Document Intelligence (prebuilt-layout)
  2. python-docx（テーブル・段落抽出フォールバック）
CLI: python pdf_to_excel.py --input <path> --output <path>
stdout: JSON {"sheets": N, "tables": N, "pages": N, "engine": "..."}
"""

import argparse
import json
import os
import re
import sys
from collections import defaultdict

import openpyxl
import pdfplumber
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# ── オプション依存ライブラリ ──────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.core.credentials import AzureKeyCredential
    HAS_DOC_INTEL = True
except ImportError:
    HAS_DOC_INTEL = False

try:
    import docx as python_docx
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


# ── 表スタイル定数 ────────────────────────────────────────────────────────
_THIN = Side(style="thin")
_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_COL_HDR_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
_COL_HDR_FONT = Font(bold=True, color="FFFFFF")
_ROW_HDR_FILL = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
_ROW_HDR_FONT = Font(bold=True, color="000000")
_STUB_FILL = PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid")
_STUB_FONT = Font(bold=True, color="000000")


def _apply_cell_style(xl_cell, kind: str | None) -> None:
    """セル種別に応じてスタイル（枠線・塗り・フォント）を適用する。"""
    xl_cell.border = _BORDER
    xl_cell.alignment = Alignment(vertical="center")
    if kind == "columnHeader":
        xl_cell.fill = _COL_HDR_FILL
        xl_cell.font = _COL_HDR_FONT
    elif kind == "rowHeader":
        xl_cell.fill = _ROW_HDR_FILL
        xl_cell.font = _ROW_HDR_FONT
    elif kind == "stub":
        xl_cell.fill = _STUB_FILL
        xl_cell.font = _STUB_FONT


# ── Doc Intelligence クライアント ─────────────────────────────────────────

def _get_doc_intel_client():
    if not HAS_DOC_INTEL:
        return None
    endpoint = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "").strip()
    key = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_KEY", "").strip()
    if not endpoint or not key:
        return None
    return DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))


# ── テキスト正規化 ────────────────────────────────────────────────────────

# コロン・セミコロン・アポストロフィもOCR誤読の千区切りとして許容する
# 先頭の - はDIが直接マイナス記号を返す場合に対応
_AMOUNT_RE = re.compile(r"^[-△▲(（]?[\d,，:;']+[)）]?$")

_NUMERIC_LINE_RE = re.compile(r"^[\d,，]+$")

# (2026/03/31) のような日付パターン — OCR誤読補正の対象外
_DATE_RE = re.compile(r"^\(?\d{4}[/／]\d{1,2}[/／]\d{1,2}\)?$")

# 数字と数字の間のスペース
_DIGIT_SPACE_RE = re.compile(r"(?<=\d) (?=\d)")


def _preprocess_ocr(text: str) -> str:
    """OCRが千区切りカンマを誤読した文字（/ ／ スペース）を補正する。
    日付パターン (YYYY/MM/DD) は変換しない。
    スペース後にカンマがある場合（241 69,530 のような別値混入）はスキップ。
    """
    t = text.strip()
    if not t or _DATE_RE.match(t):
        return t
    # / ／ をカンマに置換（例: 1412/755 → 1412,755）
    t = t.replace("/", ",").replace("／", ",")
    # 数字間スペースを除去（例: 1573 760 → 1573760）
    # スペース後にカンマがある場合はスキップ（別値混入の可能性）
    if " " in t:
        after_last_space = t[t.rfind(" ") + 1:]
        if "," not in after_last_space and "，" not in after_last_space:
            t = _DIGIT_SPACE_RE.sub("", t)
    return t


def _normalize_amount(text: str) -> str:
    """△1,234 / (1,234) / -1,234 → -1234 に正規化する。数字でなければ原文を返す。"""
    t = text.strip()
    if not t:
        return t
    negative = t.startswith(("△", "▲", "-")) or (
        (t.startswith("(") and t.endswith(")")) or
        (t.startswith("（") and t.endswith("）"))
    )
    cleaned = re.sub(r"[△▲(（)）,，:;'\s-]", "", t)
    if not cleaned.isdigit():
        return text
    val = int(cleaned)
    return str(-val if negative else val)


def _to_cell_value(text: str) -> int | float | str:
    """数値文字列を int / float に変換してExcelが数値セルとして認識できるようにする。"""
    if not text:
        return text
    try:
        return int(text)
    except ValueError:
        pass
    try:
        return float(text)
    except ValueError:
        pass
    return text


def _join_vertical_text(text: str) -> str:
    """TKC縦書きPDF対策: 改行で1文字ずつ分断されたセル値を結合する。
    例: '現\\n金\\n及\\nび\\n預\\n金' → '現金及び預金'
    各行が3文字以下かつ数字のみでない場合に限り結合する（金額行は対象外）。
    """
    if "\n" not in text:
        return text
    lines = [l.strip() for l in text.split("\n")]
    non_empty = [l for l in lines if l]
    if not non_empty:
        return text
    if all(1 <= len(l) <= 3 and not _NUMERIC_LINE_RE.match(l) for l in non_empty):
        joined = "".join(non_empty)
        joined = joined.replace("、", "").replace("・", "").replace("·", "")
        return joined
    return text


def _maybe_normalize(text: str) -> str:
    """縦書き結合 → :unselected: 除去 → OCR誤読補正 → 金額正規化の順で処理する。"""
    t = text.strip()
    # DIチェックボックスマーカーを除去（文中に混入する場合も含む）
    t = re.sub(r"\s*:(?:un)?selected:\s*", "", t).strip()
    if not t:
        return ""
    t = _join_vertical_text(t)
    t = _preprocess_ocr(t)
    if _AMOUNT_RE.match(t):
        return _normalize_amount(t)
    return t


# ── DI セル座標ユーティリティ ─────────────────────────────────────────────

def _point_xy(point):
    if isinstance(point, dict):
        return point.get("x"), point.get("y")
    if isinstance(point, (list, tuple)) and len(point) >= 2:
        return point[0], point[1]
    return getattr(point, "x", None), getattr(point, "y", None)


def _cell_x_bounds(di_cell) -> tuple[float, float] | None:
    """DI セルの左端・右端X座標を返す。"""
    regions = getattr(di_cell, "bounding_regions", None) or []
    xs: list[float] = []
    for region in regions:
        polygon = getattr(region, "polygon", None) or []
        if polygon and all(isinstance(p, (int, float)) for p in polygon):
            xs.extend(float(polygon[i]) for i in range(0, len(polygon), 2))
            continue
        for point in polygon:
            x, _y = _point_xy(point)
            if x is not None:
                xs.append(float(x))
    if not xs:
        return None
    return min(xs), max(xs)


def _get_y_from_polygon(polygon) -> float | None:
    """polygon の最小Y座標を返す（フラットリスト・ポイントリスト両対応）。"""
    ys: list[float] = []
    if polygon and all(isinstance(p, (int, float)) for p in polygon):
        ys = [float(polygon[i]) for i in range(1, len(polygon), 2)]
    else:
        for point in polygon:
            _x, y = _point_xy(point)
            if y is not None:
                ys.append(float(y))
    return min(ys) if ys else None


# ── 列マージ判定 ──────────────────────────────────────────────────────────

def _is_amountish_fragment(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    stripped = re.sub(r"[\s,，:;'\.\-\+\(\)（）△▲]", "", t)
    return bool(stripped) and stripped.isdigit()


def _looks_like_fragment(text: str) -> bool:
    """カンマ等の区切り文字で終わる場合のみ「途中で切れた数字」と見なす。
    短い独立した数値（100, 200 など）は断片扱いしない。
    """
    t = text.strip()
    if not t:
        return False
    return t.endswith((",", "，", ":", ";"))


def _join_cell_text(left: str, right: str) -> str:
    if not left:
        return right
    if not right:
        return left
    if _is_amountish_fragment(left) and _is_amountish_fragment(right):
        return _normalize_amount(left + right)
    return left + right


def _stripped_digit_len(text: str) -> int:
    """区切り文字を除去した後の桁数を返す。"""
    return len(re.sub(r"[\s,，:;'\.\-\+\(\)（）△▲]", "", text.strip()))


def _repair_extra_cells(
    grid: dict[tuple[int, int], tuple[str, str, int, int]],
) -> dict[tuple[int, int], tuple[str, str, int, int]]:
    """行のセル数が最頻値+1の行を修正する。
    「3桁以下の数値断片」と隣接する数値を結合し、以降のセルを左シフトして列数を揃える。
    ちょうど1セル多い行のみ対象（複数ズレは対象外）。
    """
    row_counts: dict[int, int] = {}
    for (r, c), (v, k, rs, cs) in grid.items():
        if v and cs == 1:
            row_counts[r] = row_counts.get(r, 0) + 1

    if not row_counts:
        return grid

    counts = list(row_counts.values())
    expected = max(set(counts), key=counts.count)

    result = dict(grid)
    for r, count in row_counts.items():
        if count != expected + 1:
            continue

        row_items = sorted(
            [(c, result[(r, c)]) for (rr, c) in result if rr == r],
            key=lambda x: x[0],
        )

        # 小数点を含む数値セルが1つでもある行は行ごとスキップ
        # （対比欄の小数値がある行でペアを誤結合するのを防ぐ）
        if any(
            "." in v and _is_amountish_fragment(v)
            for _, (v, _, _, _) in row_items
        ):
            continue

        for i in range(len(row_items) - 1):
            c1, (v1, k1, rs1, cs1) = row_items[i]
            c2, (v2, k2, rs2, cs2) = row_items[i + 1]
            if c2 != c1 + 1:
                continue
            if not (_is_amountish_fragment(v1) and _is_amountish_fragment(v2)):
                continue
            # 両方とも4桁以上の場合は独立した数値として扱い結合しない
            if _stripped_digit_len(v1) > 3 and _stripped_digit_len(v2) > 3:
                continue

            merged = _join_cell_text(v1, v2)
            del result[(r, c1)]
            del result[(r, c2)]
            result[(r, c1)] = (merged, k1, rs1, cs1)

            to_shift = sorted(
                [(c, data) for (rr, c), data in list(result.items())
                 if rr == r and c > c2],
                key=lambda x: x[0],
            )
            for c_old, data in to_shift:
                del result[(r, c_old)]
                result[(r, c_old - 1)] = data
            break

    return result


def _build_logical_column_map(raw_cells: list[dict], column_count: int) -> dict[int, int]:
    """DI の物理列番号 → 論理列番号のマッピングを返す。
    隣接する列が「途中で切れた数字（断片）」を持つ場合のみ結合する。
    財務表の独立した数値列は結合しない。
    """
    col_bounds: dict[int, list[tuple[float, float]]] = defaultdict(list)
    col_values: dict[int, list[str]] = defaultdict(list)
    row_values: dict[tuple[int, int], str] = {}
    all_xs: list[float] = []

    for cell in raw_cells:
        col = int(cell["col"])
        text = str(cell.get("content") or "").strip()
        if text:
            col_values[col].append(text)
            row_values[(int(cell["row"]), col)] = text
        bounds = cell.get("bounds")
        if bounds and cell.get("cs", 1) == 1:
            x0, x1 = bounds
            col_bounds[col].append((x0, x1))
            all_xs.extend([x0, x1])

    parents = list(range(max(column_count, 1)))

    def find(x: int) -> int:
        while parents[x] != x:
            parents[x] = parents[parents[x]]
            x = parents[x]
        return x

    def union(a: int, b: int) -> None:
        ra, rb = find(a), find(b)
        if ra != rb:
            parents[rb] = ra

    def col_span(col: int) -> tuple[float, float] | None:
        spans = col_bounds.get(col) or []
        if not spans:
            return None
        return min(x0 for x0, _ in spans), max(x1 for _, x1 in spans)

    if not all_xs:
        return {c: c for c in range(column_count)}

    table_width = max(all_xs) - min(all_xs)
    if table_width <= 0:
        return {c: c for c in range(column_count)}

    for col in range(column_count - 1):
        left_span = col_span(col)
        right_span = col_span(col + 1)
        if not left_span or not right_span:
            continue

        left_x0, left_x1 = left_span
        right_x0, right_x1 = right_span
        gap = max(0.0, right_x0 - left_x1)
        left_width = max(0.0, left_x1 - left_x0)
        right_width = max(0.0, right_x1 - right_x0)

        same_rows = set(r for r, c in row_values if c == col) & set(
            r for r, c in row_values if c == col + 1
        )
        amount_split_rows = [
            r for r in same_rows
            if _is_amountish_fragment(row_values[(r, col)] + row_values[(r, col + 1)])
        ]
        amount_split_ratio = len(amount_split_rows) / max(len(same_rows), 1)

        # 少なくとも一方が「途中で切れた数字」の見た目である場合のみ結合
        has_fragment = (
            any(_looks_like_fragment(row_values.get((r, col), "")) for r in same_rows) or
            any(_looks_like_fragment(row_values.get((r, col + 1), "")) for r in same_rows)
        )

        very_close = gap <= table_width * 0.005
        close = gap <= table_width * 0.015
        narrow_side = min(left_width, right_width) <= table_width * 0.025
        amount_split_with_fragment = bool(same_rows) and amount_split_ratio >= 0.5 and has_fragment

        if (very_close and narrow_side and has_fragment) or (close and amount_split_with_fragment):
            union(col, col + 1)

    root_to_logical: dict[int, int] = {}
    col_map: dict[int, int] = {}
    for col in range(column_count):
        root = find(col)
        if root not in root_to_logical:
            root_to_logical[root] = len(root_to_logical)
        col_map[col] = root_to_logical[root]
    return col_map


# ── Doc Intelligence メイン処理 ───────────────────────────────────────────

def _process_with_doc_intel(pdf_path: str, wb: openpyxl.Workbook) -> dict | None:
    client = _get_doc_intel_client()
    if client is None:
        print("[pdf_to_excel] Doc Intelligence: not configured, skipping.", file=sys.stderr)
        return None

    input_ext = os.path.splitext(pdf_path)[1].lower()
    content_type_map = {
        ".pdf":  "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    content_type = content_type_map.get(input_ext, "application/octet-stream")
    print(f"[pdf_to_excel] Using Azure Document Intelligence (prebuilt-layout) ext={input_ext}.", file=sys.stderr)
    try:
        with open(pdf_path, "rb") as f:
            poller = client.begin_analyze_document(
                "prebuilt-layout",
                body=f,
                content_type=content_type,
                locale="ja-JP",
            )
        result = poller.result()
    except Exception as e:
        print(f"[pdf_to_excel] Doc Intelligence failed: {e}", file=sys.stderr)
        return None

    total_pages = len(result.pages) if result.pages else 0
    total_tables = 0

    # ページごとのパラグラフを収集（表上部のタイトル・日付行を転記するため）
    paragraphs_by_page: dict[int, list[tuple[float, str]]] = defaultdict(list)
    if result.paragraphs:
        for para in result.paragraphs:
            content = (para.content or "").strip()
            if not content:
                continue
            regions = getattr(para, "bounding_regions", None) or []
            page_num_p = regions[0].page_number if regions else 1
            y = 999.0
            if regions:
                poly = getattr(regions[0], "polygon", None) or []
                y_val = _get_y_from_polygon(poly)
                if y_val is not None:
                    y = y_val
            paragraphs_by_page[page_num_p].append((y, content))
    for pn in paragraphs_by_page:
        paragraphs_by_page[pn].sort(key=lambda x: x[0])

    if result.tables:
        tables_by_page: dict[int, list] = defaultdict(list)
        for table in result.tables:
            page_num = (
                table.bounding_regions[0].page_number
                if table.bounding_regions else 1
            )
            tables_by_page[page_num].append(table)

        # 各ページで「前のテーブルの下端Y」を追跡してヘッダー行の重複を防ぐ
        page_next_table_y: dict[int, float] = {}

        for page_num in sorted(tables_by_page):
            page_tables = tables_by_page[page_num]
            for t_idx, table in enumerate(page_tables):
                total_tables += 1
                sheet_name = (
                    f"P{page_num}"
                    if len(page_tables) == 1
                    else f"P{page_num}-T{t_idx + 1}"
                )
                ws = wb.create_sheet(title=sheet_name[:31])

                # テーブル上端Yを取得
                table_top_y = 999.0
                table_regions = getattr(table, "bounding_regions", None) or []
                if table_regions:
                    poly = getattr(table_regions[0], "polygon", None) or []
                    y_val = _get_y_from_polygon(poly)
                    if y_val is not None:
                        table_top_y = y_val

                # テーブルより上にあるパラグラフ（タイトル・日付等）を先頭行に書く
                prev_y = page_next_table_y.get(page_num, 0.0)
                header_rows = 0
                for y, text in paragraphs_by_page.get(page_num, []):
                    if prev_y <= y < table_top_y:
                        ws.cell(row=header_rows + 1, column=1, value=text)
                        header_rows += 1
                page_next_table_y[page_num] = table_top_y
                row_offset = header_rows + (1 if header_rows > 0 else 0)

                # グリッド構築: (row_index, col_index) → (content, kind, row_span, col_span)
                raw_cells: list[dict] = []
                for di_cell in table.cells:
                    raw_cells.append({
                        "row": di_cell.row_index,
                        "col": di_cell.column_index,
                        "content": _maybe_normalize(di_cell.content or ""),
                        "kind": di_cell.kind or "body",
                        "rs": max(1, di_cell.row_span or 1),
                        "cs": max(1, di_cell.column_span or 1),
                        "bounds": _cell_x_bounds(di_cell),
                    })

                col_map = _build_logical_column_map(raw_cells, table.column_count)

                grid: dict[tuple[int, int], tuple[str, str, int, int]] = {}
                for raw_cell in raw_cells:
                    row = int(raw_cell["row"])
                    source_col = int(raw_cell["col"])
                    source_col_end = source_col + int(raw_cell["cs"]) - 1
                    logical_col = col_map.get(source_col, source_col)
                    logical_col_end = col_map.get(source_col_end, logical_col)
                    content = str(raw_cell["content"])
                    kind = str(raw_cell["kind"])
                    rs = int(raw_cell["rs"])
                    cs = max(1, logical_col_end - logical_col + 1)
                    key = (row, logical_col)

                    if key in grid:
                        old_content, old_kind, old_rs, old_cs = grid[key]
                        content = _join_cell_text(old_content, content)
                        if old_kind != "body":
                            kind = old_kind
                        grid[key] = (content, kind, max(old_rs, rs), max(old_cs, cs))
                    else:
                        grid[key] = (content, kind, rs, cs)

                grid = _repair_extra_cells(grid)

                # ① データ書き込み & スタイル適用
                for (r, c), (content, kind, rs, cs) in grid.items():
                    xl_cell = ws.cell(row=r + 1 + row_offset, column=c + 1, value=_to_cell_value(content))
                    _apply_cell_style(xl_cell, kind)
                    if isinstance(xl_cell.value, int):
                        xl_cell.number_format = "#,##0"
                    elif isinstance(xl_cell.value, float):
                        xl_cell.number_format = "#,##0.##"

                # ② セル結合
                for (r, c), (content, kind, rs, cs) in grid.items():
                    if rs > 1 or cs > 1:
                        ws.merge_cells(
                            start_row=r + 1 + row_offset, start_column=c + 1,
                            end_row=r + rs + row_offset, end_column=c + cs,
                        )

                # ③ 列幅自動調整（1列スパンのセルのみ参照）
                col_max: dict[int, int] = {}
                for (r, c), (content, kind, rs, cs) in grid.items():
                    if content and cs == 1:
                        w = sum(2 if ord(ch) > 127 else 1 for ch in content)
                        col_max[c] = max(col_max.get(c, 0), w)
                for c_idx, width in col_max.items():
                    col_letter = get_column_letter(c_idx + 1)
                    ws.column_dimensions[col_letter].width = min(max(width + 2, 6), 40)

                print(
                    f"[pdf_to_excel] P{page_num}-T{t_idx+1}: "
                    f"{table.row_count}rows x {table.column_count}cols "
                    f"(logical_cols={len(set(col_map.values()))})",
                    file=sys.stderr,
                )

    # テーブルが検出されなかった場合はパラグラフをテキストシートへ
    if total_tables == 0 and result.paragraphs:
        ws = wb.create_sheet(title="Text")
        for para in result.paragraphs:
            if para.content and para.content.strip():
                ws.append([para.content.strip()])
        print(
            f"[pdf_to_excel] No tables found; wrote {len(result.paragraphs)} paragraphs to 'Text'.",
            file=sys.stderr,
        )

    para_count = len(result.paragraphs) if result.paragraphs else 0
    return {
        "tables": total_tables,
        "pages": total_pages,
        "paragraphs": para_count,
        "engine": "doc_intelligence",
    }


# ── LibreOffice: DOCX → PDF 変換（EMF 画像型 Word 対策） ────────────────

def _docx_to_pdf_with_libreoffice(docx_path: str) -> str | None:
    """LibreOffice --headless で .docx を .pdf に変換して返す。失敗時は None。"""
    import subprocess
    import tempfile
    out_dir = tempfile.mkdtemp()
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            capture_output=True,
            text=True,
            timeout=120,
        )
        print(f"[pdf_to_excel] LibreOffice rc={result.returncode} {result.stderr[:200]}", file=sys.stderr)
        if result.returncode != 0:
            return None
        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(out_dir, base + ".pdf")
        return pdf_path if os.path.exists(pdf_path) else None
    except FileNotFoundError:
        print("[pdf_to_excel] libreoffice not found.", file=sys.stderr)
        return None
    except subprocess.TimeoutExpired:
        print("[pdf_to_excel] LibreOffice conversion timed out.", file=sys.stderr)
        return None


# ── python-docx フォールバック（Word専用） ───────────────────────────────

def _process_with_python_docx(docx_path: str, wb: openpyxl.Workbook) -> dict:
    """python-docx で .docx のテーブルと段落をExcelに変換する。"""
    if not HAS_PYTHON_DOCX:
        print("[pdf_to_excel] python-docx not available.", file=sys.stderr)
        return {"tables": 0, "pages": 1, "engine": "none"}

    print("[pdf_to_excel] Falling back to python-docx.", file=sys.stderr)
    doc = python_docx.Document(docx_path)
    total_tables = 0
    text_rows: list[list[str]] = []

    for t_idx, table in enumerate(doc.tables):
        total_tables += 1
        sheet_name = f"Table{t_idx + 1}"
        ws = wb.create_sheet(title=sheet_name[:31])
        for row in table.rows:
            ws.append([cell.text.strip() for cell in row.cells])
        print(f"[pdf_to_excel] python-docx table {t_idx+1}: {len(table.rows)} rows", file=sys.stderr)

    if total_tables == 0:
        for para in doc.paragraphs:
            if para.text.strip():
                text_rows.append([para.text.strip()])
        if text_rows:
            ws = wb.create_sheet(title="Text")
            for row in text_rows:
                ws.append(row)

    return {"tables": total_tables, "pages": 1, "engine": "python_docx"}


# ── pdfplumber テキストフォールバック ────────────────────────────────────

def _pdfplumber_text_lines(page) -> list[list[str]]:
    words = page.extract_words(keep_blank_chars=True, x_tolerance=4, y_tolerance=4)
    if not words:
        raw = page.extract_text() or ""
        return [[ln] for ln in raw.splitlines() if ln.strip()]
    rows_by_y: dict[int, list] = defaultdict(list)
    for w in words:
        y_key = round(float(w.get("top", 0)) / 5) * 5
        rows_by_y[y_key].append(w)
    lines: list[list[str]] = []
    for y_key in sorted(rows_by_y.keys()):
        row = sorted(rows_by_y[y_key], key=lambda w: float(w.get("x0", 0)))
        line = " ".join(w["text"] for w in row).strip()
        if line:
            lines.append([line])
    return lines


# ── PyMuPDF テキストフォールバック ──────────────────────────────────────

def _pymupdf_text_lines(pdf_path: str, page_index: int) -> list[list[str]]:
    doc = fitz.open(pdf_path)
    page = doc[page_index]
    blocks = page.get_text("blocks")
    doc.close()
    lines: list[list[str]] = []
    for block in sorted(blocks, key=lambda b: (b[1], b[0])):
        block_text = str(block[4]).strip()
        for line in block_text.splitlines():
            line = line.strip()
            if line:
                lines.append([line])
    return lines


# ── GPT-4V による精度向上（リファインモード） ────────────────────────────
# 元PDFは使わない。変換後ExcelシートをLibreOffice→PDFで画像化してVisionに渡す。

def _find_libreoffice() -> str | None:
    """Linux / Windows どちらでも LibreOffice 実行ファイルを探す。"""
    import shutil as _shutil
    candidates = [
        "libreoffice",                                           # Linux (PATH)
        "soffice",                                               # Linux 別名
        r"C:\Program Files\LibreOffice\program\soffice.exe",    # Windows 標準
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if _shutil.which(c) or os.path.isfile(c):
            return c
    return None


def _render_excel_sheet_as_pages(excel_path: str, sheet_name: str, dpi: int = 200) -> list[bytes]:
    """指定シートだけの一時xlsxを作成→LibreOffice→PDF→ページごとにPNG化して返す。
    1シートが複数PDFページになる場合もページ単位のリストで返す（縦結合しない）。"""
    import subprocess
    import tempfile
    import shutil

    if not HAS_PYMUPDF:
        print("[pdf_to_excel] PyMuPDF not available, cannot render Excel sheet", file=sys.stderr)
        return []

    lo_exe = _find_libreoffice()
    if lo_exe is None:
        print("[pdf_to_excel] LibreOffice not found, cannot render Excel sheet", file=sys.stderr)
        return []

    out_dir = tempfile.mkdtemp()
    try:
        # 1. 指定シートだけを含む一時xlsxを作成
        wb_src = openpyxl.load_workbook(excel_path)
        if sheet_name not in wb_src.sheetnames:
            print(f"[pdf_to_excel] Sheet '{sheet_name}' not found in Excel", file=sys.stderr)
            return []
        for sn in [s for s in wb_src.sheetnames if s != sheet_name]:
            del wb_src[sn]
        single_xlsx = os.path.join(out_dir, "single_sheet.xlsx")
        wb_src.save(single_xlsx)
        wb_src.close()

        # 2. LibreOffice で1シートxlsx → PDF
        import platform
        lo_cmd = [lo_exe, "--headless"]
        # Linuxコンテナではホームディレクトリが書き込めずサイレント失敗するため
        # プロセス固有のtmpディレクトリをユーザープロファイルに指定する
        if platform.system() != "Windows":
            lo_profile = f"/tmp/lo_profile_{os.getpid()}"
            lo_cmd += [f"-env:UserInstallation=file://{lo_profile}"]
        lo_cmd += ["--convert-to", "pdf", "--outdir", out_dir, single_xlsx]

        result = subprocess.run(lo_cmd, capture_output=True, text=True, timeout=60)
        print(f"[pdf_to_excel] LibreOffice returncode={result.returncode} stderr={result.stderr[:200]!r}", file=sys.stderr)
        if result.returncode != 0:
            print(f"[pdf_to_excel] LibreOffice failed (rc={result.returncode})", file=sys.stderr)
            return []

        pdf_path = os.path.join(out_dir, "single_sheet.pdf")
        out_dir_files = os.listdir(out_dir)
        print(f"[pdf_to_excel] out_dir contents after LibreOffice: {out_dir_files}", file=sys.stderr)
        if not os.path.exists(pdf_path):
            print(f"[pdf_to_excel] LibreOffice output PDF not found (expected: {pdf_path})", file=sys.stderr)
            return []

        # 3. 各ページを個別PNGとして返す（縦結合しない）
        doc = fitz.open(pdf_path)
        n_pages = len(doc)
        if n_pages == 0:
            doc.close()
            return []

        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pages_png = [doc[i].get_pixmap(matrix=mat).tobytes("png") for i in range(n_pages)]
        doc.close()

        total_bytes = sum(len(p) for p in pages_png)
        print(f"[pdf_to_excel] Rendered sheet '{sheet_name}' ({n_pages} PDF page(s)) as {n_pages} PNG(s) ({total_bytes} bytes total)", file=sys.stderr)
        return pages_png

    except subprocess.TimeoutExpired:
        print("[pdf_to_excel] LibreOffice timeout when rendering Excel sheet", file=sys.stderr)
        return []
    except Exception as e:
        print(f"[pdf_to_excel] Excel sheet render error: {e}", file=sys.stderr)
        return []
    finally:
        shutil.rmtree(out_dir, ignore_errors=True)


def _call_vision_for_table(
    page_pngs: list[bytes],
    sheet_name: str,
    existing_data: list[list[str]] | None = None,
) -> list[list[str]]:
    """GPT-4V でExcelシートの各ページ画像（複数可）からテーブルを再抽出する。
    existing_data（全行×全列の現在値）を渡すと、構造を固定したまま文字化け・数値誤りのみ修正するモードになる。
    複数ページは1リクエスト内で別々の image_url ブロックとして渡す。"""
    import base64
    import json as _json
    import urllib.request
    import urllib.error

    key = os.environ.get("AZURE_OPENAI_VISION_API_KEY", "")
    instance = os.environ.get("AZURE_OPENAI_VISION_API_INSTANCE_NAME", "")
    deployment = os.environ.get("AZURE_OPENAI_VISION_API_DEPLOYMENT_NAME", "")
    api_version = os.environ.get("AZURE_OPENAI_VISION_API_VERSION", "2024-12-01-preview")

    if not (key and instance and deployment):
        print(f"[pdf_to_excel] Vision API env not configured, skipping sheet '{sheet_name}'", file=sys.stderr)
        return []

    if not page_pngs:
        return []

    url = (
        f"https://{instance}.openai.azure.com/openai/deployments/{deployment}"
        f"/chat/completions?api-version={api_version}"
    )

    n_pages = len(page_pngs)
    if n_pages > 1:
        page_note = f"このシートは{n_pages}ページにわたっています。全ページの表データを結合して1つのJSONにまとめてください。\n"
    else:
        page_note = ""

    # 既存データが渡された場合：構造固定・値修正モード
    if existing_data:
        n_rows = len(existing_data)
        n_cols = max(len(r) for r in existing_data) if existing_data else 0
        existing_json = _json.dumps(existing_data, ensure_ascii=False)
        structure_note = (
            f"【修正モード】以下が現在のExcelシート「{sheet_name}」の内容です（{n_rows}行 × {n_cols}列）:\n"
            f"{existing_json}\n\n"
            "画像を見て、上記データの文字化け・数値の読み誤りのみ修正してください。\n"
            "【絶対禁止】行数・列数の変更 / 列の並び替え / 構造の変更\n"
            f"返すJSONのrowsは必ず{n_rows}行 × {n_cols}列を維持すること。\n"
            "修正が不要なセルもすべて含めること。\n"
        )
        text_prompt = (
            f"{page_note}"
            f"{structure_note}"
            "以下のJSON形式のみで返してください。説明文は不要。\n"
            '{"rows": [["値1", "値2", ...], ...]}\n\n'
            "数値の読み取り注意:\n"
            "- △▲や括弧で囲まれた数値はマイナス（例: △1,234 → -1234）\n"
            "- 千区切りカンマを含む数値はそのまま保持（例: 1,234,567）\n"
            "- 空白セルは空文字列\n"
            "- JSONのみ返すこと"
        )
    else:
        # フォールバック：既存データなし → 従来通りゼロから抽出
        text_prompt = (
            f"{page_note}"
            "画像はExcelシートの表です。表のデータを正確に読み取り、以下のJSON形式のみで返してください。\n"
            '{"rows": [["ヘッダー1", "ヘッダー2"], ["値1", "値2"], ...]}\n\n'
            "注意:\n"
            "- △▲や括弧で囲まれた数値はマイナス（例: △1,234 → -1234）\n"
            "- 千区切りカンマを含む数値はそのまま保持（例: 1,234,567）\n"
            "- 空白セルは空文字列\n"
            "- セル結合がある場合は各セルに同じ値を入れる\n"
            "- 複数ページの場合はヘッダー行は先頭ページのみ使用し、データ行を連結する\n"
            "- 説明文は不要。JSONのみ返すこと"
        )

    # existing_data モードは画像テキスト確認が目的なので auto でトークン節約
    # ゼロ抽出モード（existing_data なし）は high で精度優先
    img_detail = "auto" if existing_data else "high"

    # 各ページを個別の image_url ブロックとして追加
    content: list[dict] = []
    for png_bytes in page_pngs:
        img_b64 = base64.b64encode(png_bytes).decode()
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{img_b64}", "detail": img_detail},
        })
    content.append({"type": "text", "text": text_prompt})

    payload = {
        "messages": [{"role": "user", "content": content}],
        "max_completion_tokens": 16000,
    }

    print(f"[pdf_to_excel] Vision request: {len(page_pngs)} image(s) detail={img_detail} sheet='{sheet_name}'", file=sys.stderr)

    data = _json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        url, data=data, headers={"api-key": key, "Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = _json.loads(resp.read().decode())
        choice = result["choices"][0]
        finish_reason = choice.get("finish_reason", "unknown")
        usage = result.get("usage", {})
        print(
            f"[pdf_to_excel] Vision response: finish_reason={finish_reason} "
            f"prompt_tokens={usage.get('prompt_tokens', '?')} completion_tokens={usage.get('completion_tokens', '?')}",
            file=sys.stderr,
        )
        message = choice.get("message", {})
        msg_keys = list(message.keys())
        raw_content = message.get("content")
        print(f"[pdf_to_excel] Vision message keys={msg_keys} content_type={type(raw_content).__name__} content_len={len(raw_content) if isinstance(raw_content, str) else raw_content}", file=sys.stderr)
        content_str = (raw_content or "").strip()
        json_match = re.search(r"\{[\s\S]*\}", content_str)
        if json_match:
            parsed = _json.loads(json_match.group())
            rows = parsed.get("rows", [])
            return [[str(cell) for cell in row] for row in rows if row]
        print(f"[pdf_to_excel] Vision: no JSON found for sheet '{sheet_name}' (raw: {content_str[:300]})", file=sys.stderr)
    except urllib.error.HTTPError as e:
        body = e.read().decode(errors="replace")[:300]
        print(f"[pdf_to_excel] Vision HTTPError {e.code} sheet '{sheet_name}': {body}", file=sys.stderr)
    except Exception as e:
        print(f"[pdf_to_excel] Vision error sheet '{sheet_name}': {e}", file=sys.stderr)

    return []


def _refine_pages_with_vision(
    excel_path: str,
    output_path: str,
    target_sheets: list[str],
) -> dict:
    """指定Excelシートを画像化→GPT-4Vで再抽出して置き換える。元PDFは使用しない。"""
    wb = openpyxl.load_workbook(excel_path)
    refined = 0
    skipped = 0

    for sheet_name in target_sheets:
        if sheet_name not in wb.sheetnames:
            print(f"[pdf_to_excel] Sheet '{sheet_name}' not found in Excel, skipping", file=sys.stderr)
            skipped += 1
            continue

        # 既存シートの全データ（全行×全列）を読み取り Vision の修正制約として渡す
        ws_src = wb[sheet_name]
        existing_data: list[list[str]] = []
        for xl_row in ws_src.iter_rows(values_only=True):
            existing_data.append([str(v) if v is not None else "" for v in xl_row])
        n_rows = len(existing_data)
        n_cols = max(len(r) for r in existing_data) if existing_data else 0
        print(f"[pdf_to_excel] Sheet '{sheet_name}' existing data: {n_rows} rows × {n_cols} cols", file=sys.stderr)

        page_pngs = _render_excel_sheet_as_pages(excel_path, sheet_name)
        if not page_pngs:
            print(f"[pdf_to_excel] Failed to render sheet '{sheet_name}', skipping", file=sys.stderr)
            skipped += 1
            continue

        rows = _call_vision_for_table(page_pngs, sheet_name, existing_data=existing_data or None)
        if not rows:
            print(f"[pdf_to_excel] Vision returned no data for sheet '{sheet_name}', skipping", file=sys.stderr)
            skipped += 1
            continue

        # 行数・列数チェック：構造が既存データと一致しない場合は適用しない
        if existing_data and n_rows > 0 and n_cols > 0:
            actual_cols = max(len(r) for r in rows) if rows else 0
            if len(rows) != n_rows or actual_cols != n_cols:
                print(
                    f"[pdf_to_excel] Vision returned wrong dimensions for '{sheet_name}': "
                    f"expected {n_rows}×{n_cols}, got {len(rows)}×{actual_cols}. Skipping to avoid structure corruption.",
                    file=sys.stderr,
                )
                skipped += 1
                continue

        # シート順序を保ちながら置き換え
        insert_idx = wb.sheetnames.index(sheet_name)
        del wb[sheet_name]
        ws = wb.create_sheet(title=sheet_name, index=insert_idx)

        for row in rows:
            ws.append(row)

        for r_idx, xl_row in enumerate(ws.iter_rows()):
            for xl_cell in xl_row:
                if r_idx == 0:
                    _apply_cell_style(xl_cell, "columnHeader")
                else:
                    xl_cell.border = _BORDER
                    xl_cell.alignment = Alignment(vertical="center")

        # 列幅を内容の最大長に合わせて自動調整（日本語は2単位、ASCII は1単位）
        from openpyxl.utils import get_column_letter
        for col_idx in range(1, ws.max_column + 1):
            col_letter = get_column_letter(col_idx)
            max_width = 10  # 最低幅
            for cell in ws[col_letter]:
                if cell.value is not None:
                    char_width = sum(2 if ord(c) > 127 else 1 for c in str(cell.value))
                    max_width = max(max_width, char_width)
            ws.column_dimensions[col_letter].width = min(max_width + 2, 60)

        refined += 1
        col_count = max(len(r) for r in rows) if rows else 0
        print(f"[pdf_to_excel] Refined '{sheet_name}': {len(rows)} rows x {col_count} cols", file=sys.stderr)

    wb.save(output_path)
    return {"refined": refined, "skipped": skipped, "engine": "vision"}


# ── メイン ────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--refine-sheets", default=None, help="再抽出シート名のカンマ区切り（例: P2,P3,P4）。指定時は --input がExcelパスになる。")
    args = parser.parse_args()

    # ── リファインモード（--inputはExcelパス、元PDFは使わない） ──────────
    if args.refine_sheets:
        target_sheets = [s.strip() for s in args.refine_sheets.split(",") if s.strip()]
        print(f"[pdf_to_excel] refine mode: excel={args.input}, sheets={target_sheets}", file=sys.stderr)
        result = _refine_pages_with_vision(args.input, args.output, target_sheets)
        print(json.dumps(result))
        return

    print(
        f"[pdf_to_excel] HAS_PYMUPDF={HAS_PYMUPDF} HAS_DOC_INTEL={HAS_DOC_INTEL}",
        file=sys.stderr,
    )

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    input_ext = os.path.splitext(args.input)[1].lower()

    # ── .docx: python-docx を優先（DI が空成功を返す EMF 画像型 Word 対策） ──
    if input_ext == ".docx":
        docx_result = _process_with_python_docx(args.input, wb)
        got_content = docx_result["tables"] > 0 or any(
            ws.max_row and ws.max_row > 0 for ws in wb.worksheets
        )
        if got_content:
            if not wb.sheetnames:
                wb.create_sheet(title="Sheet1")
            wb.save(str(args.output))
            print(json.dumps({
                "sheets": len(wb.sheetnames),
                "tables": docx_result["tables"],
                "pages": docx_result["pages"],
                "engine": docx_result["engine"],
            }))
            return

        # python-docx で取れなかった → DI で再試行
        print("[pdf_to_excel] python-docx found nothing, trying Doc Intelligence.", file=sys.stderr)
        wb2 = openpyxl.Workbook()
        wb2.remove(wb2.active)
        doc_intel_result = _process_with_doc_intel(args.input, wb2)
        if doc_intel_result is not None and (doc_intel_result["tables"] > 0 or doc_intel_result.get("paragraphs", 0) > 0):
            if not wb2.sheetnames:
                wb2.create_sheet(title="Sheet1")
            wb2.save(str(args.output))
            print(json.dumps({
                "sheets": len(wb2.sheetnames),
                "tables": doc_intel_result["tables"],
                "pages": doc_intel_result["pages"],
                "engine": "doc_intelligence",
            }))
            return

        # python-docx も DI も空 → LibreOffice で DOCX→PDF に変換して再度 DI を試みる
        print("[pdf_to_excel] Trying LibreOffice DOCX→PDF conversion.", file=sys.stderr)
        pdf_path = _docx_to_pdf_with_libreoffice(args.input)
        if pdf_path:
            wb3 = openpyxl.Workbook()
            wb3.remove(wb3.active)
            di_pdf_result = _process_with_doc_intel(pdf_path, wb3)
            if di_pdf_result is not None and (di_pdf_result["tables"] > 0 or di_pdf_result.get("paragraphs", 0) > 0):
                if not wb3.sheetnames:
                    wb3.create_sheet(title="Sheet1")
                wb3.save(str(args.output))
                print(json.dumps({
                    "sheets": len(wb3.sheetnames),
                    "tables": di_pdf_result["tables"],
                    "pages": di_pdf_result["pages"],
                    "engine": "doc_intelligence_via_pdf",
                }))
                return

        print("[pdf_to_excel] All extraction methods failed for DOCX.", file=sys.stderr)
        ws_msg = wb.create_sheet(title="注意")
        ws_msg.append(["このWordファイルは画像埋め込み型のため、表データを抽出できませんでした。"])
        ws_msg.append(["PDFとして保存してからアップロードしてください。"])
        wb.save(str(args.output))
        print(json.dumps({"sheets": 1, "tables": 0, "pages": 1, "engine": "none"}))
        return

    # ── 1. Azure Document Intelligence（PDF 等） ─────────────────────────
    doc_intel_result = _process_with_doc_intel(args.input, wb)

    if doc_intel_result is not None:
        if not wb.sheetnames:
            wb.create_sheet(title="Sheet1")
        wb.save(str(args.output))
        print(json.dumps({
            "sheets": len(wb.sheetnames),
            "tables": doc_intel_result["tables"],
            "pages": doc_intel_result["pages"],
            "engine": "doc_intelligence",
        }))
        return

    # ── 2. フォールバック（PDF） ──────────────────────────────────────────

    if input_ext not in (".pdf",):
        print(f"[pdf_to_excel] No fallback for {input_ext}.", file=sys.stderr)
        wb.create_sheet(title="Sheet1")
        wb.save(str(args.output))
        print(json.dumps({"sheets": 1, "tables": 0, "pages": 0, "engine": "none"}))
        return

    print("[pdf_to_excel] Falling back to pdfplumber/PyMuPDF.", file=sys.stderr)
    total_tables = 0
    text_rows: list[list[str]] = []

    with pdfplumber.open(args.input) as pdf:
        total_pages = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables() or []
            if tables:
                for t_idx, table in enumerate(tables):
                    total_tables += 1
                    sheet_name = (
                        f"P{page_num}" if len(tables) == 1 else f"P{page_num}-T{t_idx + 1}"
                    )
                    ws = wb.create_sheet(title=sheet_name[:31])
                    for row in table:
                        ws.append([str(cell) if cell is not None else "" for cell in row])
                continue

            lines: list[list[str]] = []

            if HAS_PYMUPDF:
                lines = _pymupdf_text_lines(args.input, page_num - 1)
                print(f"[pdf_to_excel] p{page_num} PyMuPDF: {len(lines)} rows", file=sys.stderr)

            if not lines:
                lines = _pdfplumber_text_lines(page)
                print(f"[pdf_to_excel] p{page_num} pdfplumber text: {len(lines)} rows", file=sys.stderr)

            text_rows.extend(lines)

    if total_tables == 0 and text_rows:
        ws = wb.create_sheet(title="Text")
        for row in text_rows:
            ws.append(row)

    if not wb.sheetnames:
        wb.create_sheet(title="Sheet1")

    wb.save(str(args.output))
    print(json.dumps({
        "sheets": len(wb.sheetnames),
        "tables": total_tables,
        "pages": total_pages,
        "engine": "pdfplumber",
        "textRows": len(text_rows),
    }))


if __name__ == "__main__":
    main()
