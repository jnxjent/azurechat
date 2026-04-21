"""
pdf_to_word.py  –  PDF → Word (.docx) 変換スクリプト
優先順位:
  1. pdf2docx (レイアウト・表・フォント保持)
  2. Azure Document Intelligence (prebuilt-layout) → python-docx (フォールバック)
CLI: python pdf_to_word.py --input <path> --output <path>
stdout: JSON {"paragraphs": N, "tables": N, "engine": "..."}
"""

import argparse
import json
import os
import re
import sys

# ── オプション依存ライブラリ ──────────────────────────────────────────────
try:
    from pdf2docx import Converter as Pdf2DocxConverter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

try:
    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.core.credentials import AzureKeyCredential
    HAS_DOC_INTEL = True
except ImportError:
    HAS_DOC_INTEL = False

try:
    import docx as python_docx
    from docx.shared import Pt
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


print(
    f"[pdf_to_word] HAS_PDF2DOCX={HAS_PDF2DOCX} HAS_DOC_INTEL={HAS_DOC_INTEL} HAS_PYTHON_DOCX={HAS_PYTHON_DOCX}",
    file=sys.stderr,
)


# ── Doc Intelligence クライアント ─────────────────────────────────────────

def _get_doc_intel_client():
    if not HAS_DOC_INTEL:
        return None
    endpoint = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "").strip()
    key = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_KEY", "").strip()
    if not endpoint or not key:
        return None
    return DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))


# ── 1. pdf2docx ──────────────────────────────────────────────────────────

def _convert_with_pdf2docx(pdf_path: str, output_path: str) -> dict | None:
    """pdf2docx でレイアウトを保持した変換。成功すれば dict, 失敗すれば None。"""
    if not HAS_PDF2DOCX:
        print("[pdf_to_word] pdf2docx not available.", file=sys.stderr)
        return None
    try:
        print("[pdf_to_word] Using pdf2docx.", file=sys.stderr)
        cv = Pdf2DocxConverter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()

        # 変換後ファイルが存在するか確認
        if not os.path.exists(output_path):
            print("[pdf_to_word] pdf2docx produced no output file.", file=sys.stderr)
            return None

        # 簡易的にパラグラフ数・テーブル数を取得
        if HAS_PYTHON_DOCX:
            doc = python_docx.Document(output_path)
            return {
                "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                "tables": len(doc.tables),
                "engine": "pdf2docx",
            }
        return {"paragraphs": 0, "tables": 0, "engine": "pdf2docx"}
    except Exception as e:
        print(f"[pdf_to_word] pdf2docx failed: {e}", file=sys.stderr)
        return None


# ── CJK文字間スペース除去 ─────────────────────────────────────────────────

def _fix_cjk_spacing(text: str) -> str:
    """TKC等PDFで文字座標が個別指定されることで生じるCJK文字間の不要スペースを除去する。"""
    # CJK統合漢字・ひらがな・カタカナ・全角記号の間にある空白を削除
    return re.sub(
        r'(?<=[\u3000-\u9fff\uff00-\uffef\u3040-\u30ff])\s+(?=[\u3000-\u9fff\uff00-\uffef\u3040-\u30ff])',
        '',
        text,
    )


# ── 2. Doc Intelligence → python-docx ────────────────────────────────────

def _convert_with_doc_intel(pdf_path: str, output_path: str) -> dict | None:
    """Azure DI で段落・表を抽出して python-docx で DOCX を生成する。"""
    if not HAS_PYTHON_DOCX:
        print("[pdf_to_word] python-docx not available.", file=sys.stderr)
        return None

    client = _get_doc_intel_client()
    if client is None:
        print("[pdf_to_word] Doc Intelligence not configured.", file=sys.stderr)
        return None

    print("[pdf_to_word] Using Azure Document Intelligence (prebuilt-layout).", file=sys.stderr)
    try:
        with open(pdf_path, "rb") as f:
            poller = client.begin_analyze_document(
                "prebuilt-layout",
                body=f,
                content_type="application/pdf",
            )
        result = poller.result()
    except Exception as e:
        print(f"[pdf_to_word] Doc Intelligence failed: {e}", file=sys.stderr)
        return None

    doc = python_docx.Document()
    total_paragraphs = 0
    total_tables = 0

    # 表を先にシートとして追加
    if result.tables:
        for table in result.tables:
            total_tables += 1
            t = doc.add_table(rows=table.row_count, cols=table.column_count)
            t.style = "Table Grid"
            for cell in table.cells:
                content = (cell.content or "").replace(":selected:", "").replace(":unselected:", "").strip()
                content = _fix_cjk_spacing(content)
                t.cell(cell.row_index, cell.column_index).text = content
            doc.add_paragraph()  # 表の後に空行

    # 段落テキスト
    if result.paragraphs:
        for para in result.paragraphs:
            text = _fix_cjk_spacing((para.content or "").strip())
            if not text:
                continue
            total_paragraphs += 1
            role = getattr(para, "role", None) or ""
            if role in ("title", "sectionHeading"):
                doc.add_heading(text, level=1 if role == "title" else 2)
            else:
                doc.add_paragraph(text)

    doc.save(output_path)
    return {
        "paragraphs": total_paragraphs,
        "tables": total_tables,
        "engine": "doc_intelligence",
    }


# ── メイン ────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument(
        "--mode",
        choices=["layout", "editable"],
        default="layout",
        help="layout: pdf2docx優先（見た目再現）/ editable: DI優先（テキスト・表抽出）",
    )
    args = parser.parse_args()

    print(f"[pdf_to_word] mode={args.mode}", file=sys.stderr)

    input_ext = os.path.splitext(args.input)[1].lower()
    if input_ext != ".pdf":
        print(f"[pdf_to_word] Unsupported input type: {input_ext}", file=sys.stderr)
        print(json.dumps({"paragraphs": 0, "tables": 0, "engine": "none", "error": f"Unsupported: {input_ext}"}))
        return

    if args.mode == "editable":
        # ── editable: DI優先 → pdf2docx フォールバック ──────────────────
        result = _convert_with_doc_intel(args.input, args.output)
        if result is not None:
            print(json.dumps(result))
            return
        result = _convert_with_pdf2docx(args.input, args.output)
        if result is not None:
            print(json.dumps(result))
            return
    else:
        # ── layout: pdf2docx優先 → DI フォールバック ────────────────────
        result = _convert_with_pdf2docx(args.input, args.output)
        if result is not None:
            print(json.dumps(result))
            return
        result = _convert_with_doc_intel(args.input, args.output)
        if result is not None:
            print(json.dumps(result))
            return

    # ── 全失敗 ───────────────────────────────────────────────────────────
    print("[pdf_to_word] All conversion methods failed.", file=sys.stderr)
    print(json.dumps({"paragraphs": 0, "tables": 0, "engine": "none"}))


if __name__ == "__main__":
    main()
