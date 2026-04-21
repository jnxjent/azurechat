"""
create_word.py  –  python-docx ベースの Word 新規作成スクリプト
CLI: python create_word.py --output <path> --plan <json_path>
stdout: JSON {"paragraphs": N, "tables": N}

plan JSON schema:
{
  "title": "文書タイトル",
  "sections": [
    {
      "heading": "見出し（省略可）",
      "level": 1,           // 見出しレベル 1〜3（省略時 1）
      "paragraphs": ["本文1", "本文2"],
      "bullets": ["箇条書き1", "箇条書き2"],
      "table": {
        "headers": ["列1", "列2"],
        "rows": [["A", "B"], ["C", "D"]]
      }
    }
  ],
  "style": {
    "fontFace": "Meiryo",
    "fontSize": 11,
    "titleFontSize": 16
  }
}
"""

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_face: str, font_size_pt: float | None = None):
    """ランのフォントを設定（英数字・日本語両方）。"""
    run.font.name = font_face
    # 日本語フォントも明示的に設定
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_face)
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def add_heading_paragraph(doc: Document, text: str, level: int, font_face: str):
    """見出し段落を追加。"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, font_face)
    return heading


def add_body_paragraph(doc: Document, text: str, font_face: str, font_size_pt: float):
    """本文段落を追加。"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        set_run_font(run, font_face, font_size_pt)
    return para


def add_bullet_paragraph(doc: Document, text: str, font_face: str, font_size_pt: float):
    """箇条書き段落を追加。"""
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        set_run_font(run, font_face, font_size_pt)
    return para


def add_table_section(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    font_face: str,
    font_size_pt: float,
) -> int:
    """テーブルを追加。追加した行数を返す。"""
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return 0

    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # ヘッダー行
    header_row = table.rows[0]
    for ci, header in enumerate(headers[:col_count]):
        cell = header_row.cells[ci]
        cell.text = str(header)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                set_run_font(run, font_face, font_size_pt)
        # ヘッダー背景色（薄いグレー）
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # データ行
    for ri, row_data in enumerate(rows):
        data_row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data[:col_count]):
            cell = data_row.cells[ci]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, font_face, font_size_pt)

    doc.add_paragraph()  # テーブル後の空行
    return len(rows)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    parser.add_argument("--plan", required=True)
    args = parser.parse_args()

    output_path = Path(args.output)
    plan_path = Path(args.plan)

    plan: dict[str, Any] = json.loads(plan_path.read_text(encoding="utf-8"))

    style = plan.get("style") or {}
    font_face = str(style.get("fontFace") or "Meiryo")
    font_size_pt = float(style.get("fontSize") or 11)
    title_font_size = float(style.get("titleFontSize") or 16)

    doc = Document()

    # ── タイトル ──────────────────────────────────
    title_text = str(plan.get("title") or "").strip()
    if title_text:
        title_para = doc.add_heading(title_text, level=0)
        for run in title_para.runs:
            set_run_font(run, font_face, title_font_size)

    total_paragraphs = 0
    total_tables = 0

    # ── セクション ────────────────────────────────
    for section in plan.get("sections") or []:
        heading_text = str(section.get("heading") or "").strip()
        level = int(section.get("level") or 1)

        if heading_text:
            add_heading_paragraph(doc, heading_text, level, font_face)

        for para_text in section.get("paragraphs") or []:
            text = str(para_text).strip()
            if text:
                add_body_paragraph(doc, text, font_face, font_size_pt)
                total_paragraphs += 1

        for bullet_text in section.get("bullets") or []:
            text = str(bullet_text).strip()
            if text:
                add_bullet_paragraph(doc, text, font_face, font_size_pt)
                total_paragraphs += 1

        table_def = section.get("table")
        if table_def and isinstance(table_def, dict):
            headers = [str(h) for h in (table_def.get("headers") or [])]
            rows = [[str(c) for c in row] for row in (table_def.get("rows") or [])]
            if headers or rows:
                add_table_section(doc, headers, rows, font_face, font_size_pt)
                total_tables += 1

    doc.save(str(output_path))
    print(json.dumps({"paragraphs": total_paragraphs, "tables": total_tables}))


if __name__ == "__main__":
    main()
